"""
Export Service for LocalAIChatBox.
Exports chat history, research reports, knowledge graph data,
and analytics as JSON, CSV, Markdown, PDF, or DOCX.

Supports structured report export with table of contents,
citations, and formatted output.
"""

import csv
import io
import json
import re
from datetime import datetime
from typing import Dict, List, Optional
from sqlalchemy.orm import Session
from app.models import Conversation, ChatSession, ResearchTask, Document, UsageLog, User

import logging
logger = logging.getLogger(__name__)


def export_chat_history(db: Session, user_id: int, session_id: str = None,
                        format: str = "json") -> Dict:
    """Export chat history as JSON, CSV, or Markdown."""
    query = db.query(Conversation).filter(Conversation.user_id == user_id)
    if session_id:
        query = query.filter(Conversation.session_id == session_id)
    conversations = query.order_by(Conversation.created_at.asc()).all()

    if format == "json":
        data = []
        for conv in conversations:
            data.append({
                "id": conv.id,
                "question": conv.question,
                "answer": conv.answer,
                "sources_used": conv.sources_used,
                "search_mode": conv.search_mode or "hybrid",
                "created_at": conv.created_at.isoformat(),
                "session_id": conv.session_id,
            })
        return {
            "content": json.dumps(data, indent=2, ensure_ascii=False),
            "filename": f"chat_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json",
            "content_type": "application/json",
        }

    elif format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["ID", "Question", "Answer", "Sources", "Search Mode", "Created At"])
        for conv in conversations:
            writer.writerow([
                conv.id, conv.question, conv.answer,
                conv.sources_used, conv.search_mode or "hybrid",
                conv.created_at.isoformat(),
            ])
        return {
            "content": output.getvalue(),
            "filename": f"chat_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.csv",
            "content_type": "text/csv",
        }

    elif format == "markdown":
        md = f"# Chat History Export\n\n"
        md += f"**User ID**: {user_id}\n"
        md += f"**Exported at**: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC\n"
        md += f"**Total messages**: {len(conversations)}\n\n---\n\n"

        for conv in conversations:
            md += f"### Q: {conv.question}\n\n"
            md += f"{conv.answer}\n\n"
            if conv.sources_used:
                md += f"*Sources: {conv.sources_used}*\n\n"
            md += f"*{conv.created_at.strftime('%Y-%m-%d %H:%M')}*\n\n---\n\n"

        return {
            "content": md,
            "filename": f"chat_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.md",
            "content_type": "text/markdown",
        }

    raise ValueError(f"Unsupported format: {format}")


def export_research_report(db: Session, task_id: str, user_id: int,
                           format: str = "markdown") -> Dict:
    """Export a research report."""
    task = db.query(ResearchTask).filter(
        ResearchTask.id == task_id,
        ResearchTask.user_id == user_id,
    ).first()

    if not task:
        return None

    if format == "markdown":
        md = f"# Research Report\n\n"
        md += f"**Query**: {task.query}\n\n"
        md += f"**Strategy**: {task.strategy}\n\n"
        md += f"**Status**: {task.status}\n\n"
        md += f"**Created**: {task.created_at.strftime('%Y-%m-%d %H:%M')} UTC\n\n"
        if task.completed_at:
            md += f"**Completed**: {task.completed_at.strftime('%Y-%m-%d %H:%M')} UTC\n\n"
        md += f"---\n\n"

        if task.result_knowledge:
            md += f"## Findings\n\n{task.result_knowledge}\n\n"

        if task.result_report:
            md += f"## Detailed Report\n\n{task.result_report}\n\n"

        if task.result_sources:
            md += f"## Sources\n\n"
            try:
                sources = json.loads(task.result_sources)
                for i, src in enumerate(sources, 1):
                    if isinstance(src, dict):
                        md += f"{i}. **{src.get('title', 'Source')}**\n"
                        if src.get('url'):
                            md += f"   URL: {src['url']}\n"
                        if src.get('snippet'):
                            md += f"   {src['snippet']}\n"
                    else:
                        md += f"{i}. {src}\n"
                    md += "\n"
            except (json.JSONDecodeError, TypeError):
                md += f"{task.result_sources}\n\n"

        return {
            "content": md,
            "filename": f"research_{task_id[:8]}_{datetime.utcnow().strftime('%Y%m%d')}.md",
            "content_type": "text/markdown",
        }

    elif format == "json":
        data = {
            "id": task.id,
            "query": task.query,
            "strategy": task.strategy,
            "status": task.status,
            "created_at": task.created_at.isoformat(),
            "completed_at": task.completed_at.isoformat() if task.completed_at else None,
            "findings": task.result_knowledge,
            "report": task.result_report,
            "sources": json.loads(task.result_sources) if task.result_sources else [],
            "metadata": json.loads(task.result_metadata) if task.result_metadata else {},
        }
        return {
            "content": json.dumps(data, indent=2, ensure_ascii=False),
            "filename": f"research_{task_id[:8]}_{datetime.utcnow().strftime('%Y%m%d')}.json",
            "content_type": "application/json",
        }

    elif format == "pdf":
        return _export_research_pdf(task, task_id)

    elif format == "docx":
        return _export_research_docx(task, task_id)

    raise ValueError(f"Unsupported format: {format}")


def _export_research_pdf(task, task_id: str) -> Dict:
    """Export research report as PDF using fpdf2."""
    try:
        from fpdf import FPDF
    except ImportError:
        logger.warning("fpdf2 not installed, falling back to markdown")
        raise ValueError("PDF export requires fpdf2 package. Install with: pip install fpdf2")

    class ResearchPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 10)
            self.cell(0, 10, "LocalAIChatBox - Research Report", border=False, align="R")
            self.ln(5)
            self.set_draw_color(66, 133, 244)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

        def chapter_title(self, title):
            self.set_font("Helvetica", "B", 14)
            self.set_text_color(33, 37, 41)
            self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(66, 133, 244)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(4)

        def chapter_body(self, text):
            self.set_font("Helvetica", "", 10)
            self.set_text_color(52, 58, 64)
            # Clean markdown formatting for PDF
            clean = _clean_markdown_for_pdf(text)
            self.multi_cell(0, 6, clean)
            self.ln(4)

    pdf = ResearchPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(33, 37, 41)
    pdf.cell(0, 15, "Research Report", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    # Metadata
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(108, 117, 125)
    pdf.cell(0, 6, f"Query: {task.query}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Strategy: {task.strategy}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Created: {task.created_at.strftime('%Y-%m-%d %H:%M')} UTC", new_x="LMARGIN", new_y="NEXT")
    if task.completed_at:
        pdf.cell(0, 6, f"Completed: {task.completed_at.strftime('%Y-%m-%d %H:%M')} UTC", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)

    # Findings
    if task.result_knowledge:
        pdf.chapter_title("Key Findings")
        pdf.chapter_body(task.result_knowledge)

    # Report
    if task.result_report:
        pdf.chapter_title("Detailed Report")
        # Split report by sections (## headers)
        sections = re.split(r'\n##\s+', task.result_report)
        for i, section in enumerate(sections):
            if i == 0:
                pdf.chapter_body(section)
            else:
                lines = section.split('\n', 1)
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(52, 58, 64)
                pdf.cell(0, 8, lines[0].strip(), new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
                if len(lines) > 1:
                    pdf.chapter_body(lines[1])

    # Sources
    if task.result_sources:
        pdf.chapter_title("Sources")
        try:
            sources = json.loads(task.result_sources)
            pdf.set_font("Helvetica", "", 9)
            for i, src in enumerate(sources, 1):
                if isinstance(src, dict):
                    title = src.get('title', 'Source')
                    url = src.get('url', '')
                    pdf.set_text_color(33, 37, 41)
                    pdf.cell(0, 5, f"[{i}] {title}", new_x="LMARGIN", new_y="NEXT")
                    if url:
                        pdf.set_text_color(66, 133, 244)
                        pdf.cell(0, 5, f"    {url}", new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.set_text_color(33, 37, 41)
                    pdf.cell(0, 5, f"[{i}] {src}", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
        except (json.JSONDecodeError, TypeError):
            pdf.chapter_body(task.result_sources)

    # Output
    pdf_bytes = pdf.output()
    import base64
    return {
        "content": base64.b64encode(pdf_bytes).decode("utf-8"),
        "filename": f"research_{task_id[:8]}_{datetime.utcnow().strftime('%Y%m%d')}.pdf",
        "content_type": "application/pdf",
        "is_binary": True,
    }


def _export_research_docx(task, task_id: str) -> Dict:
    """Export research report as DOCX."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ValueError("DOCX export requires python-docx package")

    doc = DocxDocument()

    # Document title
    title = doc.add_heading("Research Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Shading Accent 1'
    cells = table.rows[0].cells
    cells[0].text = "Query"
    cells[1].text = task.query
    cells = table.rows[1].cells
    cells[0].text = "Strategy"
    cells[1].text = task.strategy
    cells = table.rows[2].cells
    cells[0].text = "Created"
    cells[1].text = task.created_at.strftime('%Y-%m-%d %H:%M') + " UTC"
    cells = table.rows[3].cells
    cells[0].text = "Status"
    cells[1].text = task.status

    doc.add_paragraph()

    # Findings
    if task.result_knowledge:
        doc.add_heading("Key Findings", level=1)
        for paragraph in task.result_knowledge.split('\n'):
            p = paragraph.strip()
            if p:
                if p.startswith('- ') or p.startswith('* '):
                    doc.add_paragraph(p[2:], style='List Bullet')
                elif re.match(r'^\d+\.', p):
                    doc.add_paragraph(re.sub(r'^\d+\.\s*', '', p), style='List Number')
                else:
                    doc.add_paragraph(p)

    # Report
    if task.result_report:
        doc.add_heading("Detailed Report", level=1)
        sections = re.split(r'\n##\s+', task.result_report)
        for i, section in enumerate(sections):
            if i == 0:
                for para in section.split('\n'):
                    p = para.strip()
                    if p:
                        _add_docx_paragraph(doc, p)
            else:
                lines = section.split('\n', 1)
                doc.add_heading(lines[0].strip(), level=2)
                if len(lines) > 1:
                    for para in lines[1].split('\n'):
                        p = para.strip()
                        if p:
                            _add_docx_paragraph(doc, p)

    # Sources
    if task.result_sources:
        doc.add_heading("Sources", level=1)
        try:
            sources = json.loads(task.result_sources)
            for i, src in enumerate(sources, 1):
                if isinstance(src, dict):
                    title_text = src.get('title', 'Source')
                    url = src.get('url', '')
                    p = doc.add_paragraph(style='List Number')
                    run = p.add_run(title_text)
                    run.bold = True
                    if url:
                        p.add_run(f"\n{url}").font.color.rgb = RGBColor(66, 133, 244)
                else:
                    doc.add_paragraph(str(src), style='List Number')
        except (json.JSONDecodeError, TypeError):
            doc.add_paragraph(task.result_sources)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)

    import base64
    return {
        "content": base64.b64encode(buffer.getvalue()).decode("utf-8"),
        "filename": f"research_{task_id[:8]}_{datetime.utcnow().strftime('%Y%m%d')}.docx",
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "is_binary": True,
    }


def _clean_markdown_for_pdf(text: str) -> str:
    """Remove markdown formatting for PDF rendering."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # italic
    text = re.sub(r'`(.+?)`', r'\1', text)  # code
    text = re.sub(r'#{1,6}\s*', '', text)  # headings
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # links
    text = re.sub(r'!\[.*?\]\(.+?\)', '', text)  # images
    return text.strip()


def _add_docx_paragraph(doc, text: str):
    """Add a paragraph to DOCX with basic markdown parsing."""
    if text.startswith('- ') or text.startswith('* '):
        doc.add_paragraph(text[2:], style='List Bullet')
    elif re.match(r'^\d+\.', text):
        doc.add_paragraph(re.sub(r'^\d+\.\s*', '', text), style='List Number')
    elif text.startswith('### '):
        doc.add_heading(text[4:], level=3)
    elif text.startswith('> '):
        p = doc.add_paragraph(text[2:])
        p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'
    else:
        doc.add_paragraph(text)


def export_knowledge_graph(kg_engine, format: str = "json") -> Dict:
    """Export knowledge graph data."""
    all_entities = kg_engine.get_all_entities()
    stats = kg_engine.get_stats()

    # Get all edges
    edges = []
    for u, v, data in kg_engine.graph.edges(data=True):
        edges.append({
            "source": u,
            "target": v,
            "relation": data.get("relation", "RELATED_TO"),
            "source_doc": data.get("source_doc", ""),
            "source_file": data.get("source_file", ""),
        })

    if format == "json":
        data = {
            "stats": stats,
            "entities": all_entities,
            "edges": edges,
            "exported_at": datetime.utcnow().isoformat(),
        }
        return {
            "content": json.dumps(data, indent=2, ensure_ascii=False),
            "filename": f"knowledge_graph_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json",
            "content_type": "application/json",
        }

    elif format == "csv":
        # Export nodes CSV
        nodes_output = io.StringIO()
        writer = csv.writer(nodes_output)
        writer.writerow(["Name", "Type", "Connections", "Source Files"])
        for entity in all_entities:
            writer.writerow([
                entity.get("name", ""),
                entity.get("type", "CONCEPT"),
                entity.get("connections", 0),
                "; ".join(entity.get("source_files", [])),
            ])

        # Export edges CSV
        edges_output = io.StringIO()
        writer2 = csv.writer(edges_output)
        writer2.writerow(["Source", "Target", "Relation", "Source File"])
        for edge in edges:
            writer2.writerow([
                edge["source"], edge["target"],
                edge["relation"], edge.get("source_file", ""),
            ])

        combined = f"=== ENTITIES ===\n{nodes_output.getvalue()}\n\n=== RELATIONSHIPS ===\n{edges_output.getvalue()}"
        return {
            "content": combined,
            "filename": f"knowledge_graph_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.csv",
            "content_type": "text/csv",
        }

    elif format == "graphml":
        # Export as GraphML for use in external graph tools
        import networkx as nx
        output = io.BytesIO()
        nx.write_graphml(kg_engine.graph, output)
        return {
            "content": output.getvalue().decode("utf-8"),
            "filename": f"knowledge_graph_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.graphml",
            "content_type": "application/xml",
        }

    raise ValueError(f"Unsupported format: {format}")


def export_documents_list(db: Session, format: str = "csv") -> Dict:
    """Export document list with metadata."""
    documents = db.query(Document).order_by(Document.uploaded_at.desc()).all()

    if format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["ID", "Filename", "Type", "Size (MB)", "Chunks",
                         "Uploaded By", "Uploaded At", "Indexed", "Folder ID", "Version"])
        for doc in documents:
            writer.writerow([
                doc.id, doc.filename, doc.file_type, doc.file_size_mb,
                doc.num_chunks, doc.uploader.username if doc.uploader else "Unknown",
                doc.uploaded_at.isoformat(), doc.is_indexed,
                doc.folder_id, doc.version,
            ])
        return {
            "content": output.getvalue(),
            "filename": f"documents_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.csv",
            "content_type": "text/csv",
        }

    elif format == "json":
        data = []
        for doc in documents:
            data.append({
                "id": doc.id,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "file_size_mb": doc.file_size_mb,
                "num_chunks": doc.num_chunks,
                "uploaded_by": doc.uploader.username if doc.uploader else "Unknown",
                "uploaded_at": doc.uploaded_at.isoformat(),
                "is_indexed": doc.is_indexed,
                "folder_id": doc.folder_id,
                "version": doc.version,
                "description": doc.description,
            })
        return {
            "content": json.dumps(data, indent=2, ensure_ascii=False),
            "filename": f"documents_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json",
            "content_type": "application/json",
        }

    raise ValueError(f"Unsupported format: {format}")
